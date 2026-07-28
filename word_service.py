import os
import io
import datetime
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=20, bottom=20, left=20, right=20):
    """Sets tight padding (20 dxa) to maximize image display area."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_table_borders(table, color_hex="D0D0D0"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="6" w:space="0" w:color="{color_hex}"/>
            <w:bottom w:val="single" w:sz="6" w:space="0" w:color="{color_hex}"/>
            <w:insideH w:val="single" w:sz="6" w:space="0" w:color="{color_hex}"/>
            <w:insideV w:val="single" w:sz="6" w:space="0" w:color="{color_hex}"/>
            <w:left w:val="single" w:sz="6" w:space="0" w:color="{color_hex}"/>
            <w:right w:val="single" w:sz="6" w:space="0" w:color="{color_hex}"/>
        </w:tblBorders>
    ''')
    tblPr.append(borders)

def get_brt_now_str():
    """Returns formatted datetime string in Brasilia timezone (UTC-3)."""
    now_utc = datetime.datetime.now(datetime.timezone.utc)
    brt_tz = datetime.timezone(datetime.timedelta(hours=-3))
    now_brt = now_utc.astimezone(brt_tz)
    return now_brt.strftime('%d/%m/%Y às %H:%M')

def fmt_num_br(val):
    """Formats numeric value or string to Brazilian currency format (1.500,00)."""
    if val is None or pd.isna(val) or str(val).strip() in ['', '-']:
        return '-'
    if isinstance(val, (int, float)):
        return f"{val:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
    s = str(val).strip().replace('R$', '').replace('(REAL)', '').replace(' ', '')
    try:
        if ',' in s and '.' in s:
            s = s.replace('.', '').replace(',', '.')
        elif ',' in s:
            s = s.replace(',', '.')
        num = float(s)
        return f"{num:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
    except Exception:
        return str(val)

def gerar_relatorio_word(resultados, df_original=None, modelo_bytes=None, titulo="Relatório de Atualização Monetária - BCB", **kwargs):
    """
    Generates Word document matching user's requested layout:
    1. Grid table (3 columns) of tightly cropped BCB screenshot prints.
    2. Table below with original columns + corrected values.
    """
    if modelo_bytes:
        doc = Document(io.BytesIO(modelo_bytes))
    else:
        doc = Document()
        
    # Page Margins (0.4 in for maximum width and readability)
    for section in doc.sections:
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.4)
        section.right_margin = Inches(0.4)
        
    # Document Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run(titulo)
    r_title.font.name = 'Calibri'
    r_title.font.size = Pt(20)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(0, 51, 102)
    
    # Date Subtitle (Brasilia local time UTC-3)
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run(f"Data de Emissão: {get_brt_now_str()} | Fonte: Banco Central do Brasil (BCB)")
    r_sub.font.name = 'Calibri'
    r_sub.font.size = Pt(9.5)
    r_sub.font.italic = True
    r_sub.font.color.rgb = RGBColor(100, 100, 100)
    
    p_sub.paragraph_format.space_after = Pt(10)

    # Section 1: Grid of Screenshots (3 per row)
    doc.add_heading("Comprovantes de Cálculo (Prints BCB)", level=2)
    
    num_items = len(resultados)
    cols_per_row = 3
    num_rows = (num_items + cols_per_row - 1) // cols_per_row if num_items > 0 else 1
    
    grid_table = doc.add_table(rows=num_rows, cols=cols_per_row)
    grid_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    grid_table.autofit = False
    set_table_borders(grid_table, "A0A0A0")
    
    for idx, r in enumerate(resultados):
        row_idx = idx // cols_per_row
        col_idx = idx % cols_per_row
        cell = grid_table.cell(row_idx, col_idx)
        set_cell_margins(cell, top=10, bottom=10, left=10, right=10)
        
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        
        img_bytes = r.get('screenshot_bytes')
        if img_bytes:
            img_stream = io.BytesIO(img_bytes)
            run = p.add_run()
            # 2.4 inches width fills the 3-column cell cleanly and yields large legible text
            run.add_picture(img_stream, width=Inches(2.4))
        else:
            p.add_run(f"[Item #{idx+1} - Sem Print]")

    for idx in range(num_items, num_rows * cols_per_row):
        row_idx = idx // cols_per_row
        col_idx = idx % cols_per_row
        cell = grid_table.cell(row_idx, col_idx)
        set_cell_background(cell, "F9F9F9")

    doc.add_paragraph().paragraph_format.space_after = Pt(14)
    
    # Section 2: Table of Updated Values Below
    doc.add_heading("Planilha de Valores Atualizados", level=2)
    
    if df_original is not None:
        table_df = df_original.copy()
        
        # Format original currency columns cleanly (1.500,00)
        for c in table_df.columns:
            c_low = str(c).lower()
            if any(k in c_low for k in ['valor', 'val', 'quantia', 'preco', 'preço', 'montante']):
                table_df[c] = table_df[c].apply(fmt_num_br)

        if len(resultados) == len(table_df):
            # Add Data da Atualização column if not present
            if 'Data da Atualização' not in table_df.columns and 'Data Final (Correção)' not in table_df.columns:
                dt_col_idx = 2 if len(table_df.columns) >= 2 else len(table_df.columns)
                table_df.insert(dt_col_idx, 'Data da Atualização', [r.get('data_final', '-') for r in resultados])
                
            table_df['Valor Corrigido (R$)'] = [fmt_num_br(r.get('valor_corrigido_str', '-')) for r in resultados]
            table_df['Fator BCB'] = [r.get('fator_correcao', '-') for r in resultados]
        
        table_res = doc.add_table(rows=1, cols=len(table_df.columns))
        table_res.alignment = WD_TABLE_ALIGNMENT.CENTER
        table_res.autofit = False
        set_table_borders(table_res, "003366")
        
        hdr_cells = table_res.rows[0].cells
        for col_i, col_name in enumerate(table_df.columns):
            hdr_cells[col_i].text = str(col_name)
            set_cell_background(hdr_cells[col_i], "003366")
            p = hdr_cells[col_i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 255, 255)
                
        for row_i, row in table_df.iterrows():
            row_cells = table_res.add_row().cells
            bg_color = "F4F6F9" if row_i % 2 == 1 else "FFFFFF"
            for col_i, val in enumerate(row):
                val_str = str(val) if not pd.isna(val) else ""
                row_cells[col_i].text = val_str
                set_cell_background(row_cells[col_i], bg_color)
                p = row_cells[col_i].paragraphs[0]
                
                col_name = str(table_df.columns[col_i]).lower()
                if any(k in col_name for k in ['valor', 'corrigido', 'fator', 'r$']):
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif any(k in col_name for k in ['data', 'dt', 'vencimento', 'status']):
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                for run in p.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(9.5)
    else:
        table_res = doc.add_table(rows=1, cols=6)
        table_res.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ["Item", "Descrição", "Data Inicial", "Data da Atualização", "Valor Original (R$)", "Valor Corrigido (R$)"]
        hdr_cells = table_res.rows[0].cells
        for i, h_text in enumerate(headers):
            hdr_cells[i].text = h_text
            set_cell_background(hdr_cells[i], "003366")
            p = hdr_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 255, 255)
                
        for idx, r in enumerate(resultados, 1):
            row_cells = table_res.add_row().cells
            vals = [
                str(idx),
                r.get('descricao', f"Item #{idx}"),
                r.get('data_inicial', ''),
                r.get('data_final', ''),
                fmt_num_br(r.get('valor_original_str', '0,00')),
                fmt_num_br(r.get('valor_corrigido_str', '-'))
            ]
            for i, val in enumerate(vals):
                row_cells[i].text = val
                p = row_cells[i].paragraphs[0]
                if i in [4, 5]:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif i in [0, 2, 3]:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(9.5)

    out_buf = io.BytesIO()
    doc.save(out_buf)
    return out_buf.getvalue()
