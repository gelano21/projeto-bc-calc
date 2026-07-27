import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def gerar_arquivos_exemplo():
    # 1. Excel Sample
    data = [
        {"Descrição": "Pagamento Fornecedor A", "Data de Pagamento": "15/01/2021", "Valor": 1500.00},
        {"Descrição": "Reembolso Despesas B", "Data de Pagamento": "10/05/2022", "Valor": 850.50},
        {"Descrição": "Parcela Contrato C", "Data de Pagamento": "01/10/2023", "Valor": 3200.00},
        {"Descrição": "Depósito Judicial D", "Data de Pagamento": "05/02/2024", "Valor": 5000.00},
    ]
    df = pd.DataFrame(data)
    df.to_excel("exemplo_pagamentos.xlsx", index=False)
    print("Criado: exemplo_pagamentos.xlsx")
    
    # 2. Word Template Sample
    doc = Document()
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run("MODELO DE RELATÓRIO DE CORREÇÃO MONETÁRIA")
    run.font.name = 'Calibri'
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    doc.add_paragraph("Este é um modelo de documento pré-formatado para inclusão automática dos cálculos do Banco Central do Brasil (BCB).")
    doc.add_paragraph()
    
    p_ph = doc.add_paragraph("{{PRINTS_CALCULO}}")
    p_ph.runs[0].font.italic = True
    p_ph.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    doc.save("modelo_exemplo.docx")
    print("Criado: modelo_exemplo.docx")

if __name__ == "__main__":
    gerar_arquivos_exemplo()
